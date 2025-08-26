"""Outils d'exportation des résultats de batch en document DOCX."""

from __future__ import annotations

import io
import json
from dataclasses import asdict, is_dataclass
from typing import Any, Dict, List, Optional

import markdown as md
from bs4 import BeautifulSoup
from bs4.element import NavigableString, Tag
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Pt, RGBColor
from docx.text.run import Run


class MarkdownToDocxConverter:
    """Convertit du texte Markdown en éléments DOCX."""

    def __init__(self, document: Document, styles: Dict[str, Dict[str, Any]]):
        """Initialise le convertisseur avec un document et un dictionnaire de styles."""

        self.doc = document
        self.styles = styles or {}

    def _apply_style(
        self,
        run,
        style_overrides: Dict[str, Any] | None = None,
        *,
        style_name: str = "response",
    ) -> None:
        """Applique un style au ``run`` donné.

        ``style_name`` permet de sélectionner un style de base dans ``self.styles``.
        ``style_overrides`` peut être utilisé pour modifier certains attributs.
        """

        style = {**self.styles.get(style_name, {}), **(style_overrides or {})}

        if font_name := style.get("font_name"):
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        if size := style.get("font_size"):
            run.font.size = Pt(size)
        if color := style.get("font_color_rgb"):
            try:
                if isinstance(color, str):
                    run.font.color.rgb = RGBColor.from_string(color)
                else:
                    run.font.color.rgb = RGBColor(*color)
            except Exception:
                pass
        run.bold = style.get("is_bold", False)
        run.italic = style.get("is_italic", False)

    def _add_inline(self, paragraph, node) -> None:
        """Ajoute récursivement les noeuds inline à un paragraphe."""

        if isinstance(node, NavigableString):
            text = str(node)
            if text:
                run = paragraph.add_run(text)
                self._apply_style(run)
            return

        for child in node.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text:
                    run = paragraph.add_run(text)
                    self._apply_style(run)
            elif child.name in {"strong", "b"}:
                run = paragraph.add_run(child.get_text())
                self._apply_style(run, {"is_bold": True})
            elif child.name in {"em", "i"}:
                run = paragraph.add_run(child.get_text())
                self._apply_style(run, {"is_italic": True})
            elif child.name == "code":
                run = paragraph.add_run(child.get_text())
                self._apply_style(run)
                run.font.name = "Consolas"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            elif child.name == "a":
                text = child.get_text()
                href = child.get("href")
                if href:
                    run = self._add_hyperlink(paragraph, href, text)
                    self._apply_style(run)
                else:
                    run = paragraph.add_run(text)
                    self._apply_style(run)
            else:
                self._add_inline(paragraph, child)

    def _add_hyperlink(self, paragraph, url: str, text: str) -> Run:
        """Ajoute un hyperlien cliquable au paragraphe."""

        part = paragraph.part
        r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = OxmlElement('w:r')
        r_pr = OxmlElement('w:rPr')
        r_style = OxmlElement('w:rStyle')
        r_style.set(qn('w:val'), 'Hyperlink')
        r_pr.append(r_style)
        new_run.append(r_pr)

        text_elem = OxmlElement('w:t')
        text_elem.text = text
        new_run.append(text_elem)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

        return Run(new_run, paragraph)

    def _process_element(self, elem, list_style: str | None = None) -> None:
        """Traite les éléments HTML convertis depuis le Markdown."""

        if isinstance(elem, NavigableString):
            text = str(elem).strip()
            if text:
                paragraph = (
                    self.doc.add_paragraph(style=list_style)
                    if list_style
                    else self.doc.add_paragraph()
                )
                run = paragraph.add_run(text)
                self._apply_style(run)
            return

        tag = elem.name
        if tag in {"p", "li"}:
            paragraph = (
                self.doc.add_paragraph(style=list_style)
                if list_style
                else self.doc.add_paragraph()
            )
            self._add_inline(paragraph, elem)
            for child in elem.find_all(["ul", "ol"], recursive=False):
                self._process_element(
                    child,
                    "List Bullet" if child.name == "ul" else "List Number",
                )
        elif tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            level = int(tag[1])
            paragraph = self.doc.add_heading(level=level)
            self._add_inline(paragraph, elem)
        elif tag == "ul":
            for li in elem.find_all("li", recursive=False):
                self._process_element(li, "List Bullet")
        elif tag == "ol":
            for li in elem.find_all("li", recursive=False):
                self._process_element(li, "List Number")
        elif tag == "pre":
            code_text = "".join(elem.strings).strip()
            paragraph = self.doc.add_paragraph()
            run = paragraph.add_run(code_text)
            self._apply_style(run)
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        elif tag == "table":
            rows = elem.find_all("tr", recursive=False)
            if rows:
                first_row_cells = rows[0].find_all(["th", "td"], recursive=False)
                cols = len(first_row_cells)
                table = self.doc.add_table(rows=len(rows), cols=cols)
                for r_idx, row in enumerate(rows):
                    cells = row.find_all(["th", "td"], recursive=False)
                    for c_idx, cell in enumerate(cells):
                        paragraph = table.cell(r_idx, c_idx).paragraphs[0]
                        self._add_inline(paragraph, cell)
        else:
            text = elem.get_text(strip=True)
            if text:
                paragraph = self.doc.add_paragraph()
                self._add_inline(paragraph, elem)

    def add_markdown(self, text: str) -> None:
        """Convertit un texte Markdown et l'ajoute au document avec un fallback."""

        try:
            if not text:
                return

            md_converter = md.Markdown(extensions=["fenced_code", "tables"])
            html = md_converter.convert(text)

            soup = BeautifulSoup(html, "lxml")
            if soup.body:
                for elem in soup.body.find_all(recursive=False):
                    self._process_element(elem)
        except Exception as e:  # pragma: no cover - fallback branch
            warning_p = self.doc.add_paragraph()
            warning_run = warning_p.add_run(
                f"[Le formatage de ce bloc a échoué. Contenu original ci-dessous. Erreur : {e}]"
            )
            warning_run.font.italic = True
            warning_run.font.color.rgb = RGBColor(120, 120, 120)

            self.doc.add_paragraph(text)

def generer_export_docx_batch(
    resultats: List[Any],
    styles_interface: Dict[str, Dict[str, Any]],
    template_source: Optional[Dict[str, Any]] = None,
) -> io.BytesIO:
    """Génère un document DOCX à partir d'une liste de résultats de batch.

    ``styles_interface`` correspond aux styles définis dans l'interface.
    ``template_source`` peut contenir des styles extraits d'un document importé
    et est prioritaire sur ``styles_interface`` lorsqu'il est fourni. Lorsque
    ``template_source`` est ``None``, les styles de l'interface servent de
    solution de repli.

    Chaque résultat doit contenir au minimum les champs ``status``,
    ``prompt_text`` et ``clean_response`` (ou ``response``).
    """

    document = Document()
    styles = template_source if template_source is not None else styles_interface
    converter = MarkdownToDocxConverter(document, styles)

    succeeded: List[Dict[str, Any]] = []
    failed: List[Dict[str, Any]] = []

    for res in resultats:
        data = asdict(res) if is_dataclass(res) else dict(res)
        if data.get("status") == "succeeded":
            succeeded.append(data)
        else:
            failed.append(data)

    # Section principale : prompts et réponses
    for item in succeeded:
        prompt_text = item.get("prompt_text", "")
        para = converter.doc.add_paragraph()
        run = para.add_run(prompt_text)
        converter._apply_style(run, style_name="prompt")

        reponse_structuree = item.get("structured_response")
        if reponse_structuree:
            for bloc in reponse_structuree:
                bloc_type = bloc.get("type", "paragraph")
                if bloc_type == "list":
                    for li in bloc.get("items", []):
                        converter.doc.add_paragraph(li, style="List Bullet")
                    continue
                if bloc_type == "table":
                    rows = bloc.get("rows", [])
                    if rows:
                        table = converter.doc.add_table(rows=len(rows), cols=len(rows[0]))
                        for r_idx, row in enumerate(rows):
                            for c_idx, cell_text in enumerate(row):
                                p = table.cell(r_idx, c_idx).paragraphs[0]
                                r = p.add_run(cell_text)
                                converter._apply_style(r)
                    continue

                if bloc_type.startswith("heading_"):
                    level = int(bloc_type.split("_")[-1])
                    p = converter.doc.add_heading(level=level)
                else:
                    p = converter.doc.add_paragraph()

                for run_data in bloc.get("runs", []):
                    r = p.add_run(run_data.get("text", ""))
                    converter._apply_style(r, run_data.get("style"), style_name="response")
        else:
            response_text = item.get("clean_response") or item.get("response", "")
            converter.add_markdown(response_text)

        converter.doc.add_paragraph()

    # Section annexe pour les erreurs
    if failed:
        converter.doc.add_page_break()
        converter.doc.add_heading("Annexe - Requêtes échouées", level=1)
        for item in failed:
            title = item.get("prompt_text") or item.get("custom_id")
            converter.doc.add_paragraph(title, style="List Bullet")
            err = item.get("error")
            if isinstance(err, dict):
                err = json.dumps(err, ensure_ascii=False, indent=2)
            converter.doc.add_paragraph(str(err))

    output = io.BytesIO()
    converter.doc.save(output)
    output.seek(0)
    return output


def _appliquer_style_run(run: Run, style: Dict[str, Any] | None, base_style: Dict[str, Any]) -> None:
    """Applique le style combiné au run."""
    style_combined = {**base_style, **(style or {})}

    font_name = style_combined.get("font_name")
    if font_name:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size := style_combined.get("font_size"):
        try:
            run.font.size = Pt(int(size))
        except Exception:
            pass
    if color := style_combined.get("font_color_rgb"):
        try:
            if isinstance(color, str):
                if color.startswith("RGBColor"):
                    parts = color[color.find("(") + 1 : color.find(")")].split(",")
                    rgb = [int(p.strip().replace("0x", ""), 16) for p in parts]
                    run.font.color.rgb = RGBColor(*rgb)
                else:
                    run.font.color.rgb = RGBColor.from_string(color)
            else:
                run.font.color.rgb = RGBColor(*color)
        except Exception:
            pass
    run.bold = style_combined.get("is_bold", False)
    run.italic = style_combined.get("is_italic", False)


def generer_export_docx(
    reponse_structuree: List[Dict[str, Any]],
    styles_interface: Dict[str, Dict[str, Any]],
) -> io.BytesIO:
    """Reconstruit un document DOCX à partir d'une structure de blocs et de runs."""

    document = Document()
    base_style = styles_interface.get("response", {})

    for bloc in reponse_structuree:
        type_bloc = bloc.get("type", "paragraph")

        if type_bloc == "list":
            for item in bloc.get("items", []):
                document.add_paragraph(item, style="List Bullet")
            continue

        if type_bloc == "table":
            rows = bloc.get("rows", [])
            if rows:
                table = document.add_table(rows=len(rows), cols=len(rows[0]))
                for r_idx, row in enumerate(rows):
                    for c_idx, cell_text in enumerate(row):
                        p = table.cell(r_idx, c_idx).paragraphs[0]
                        run = p.add_run(cell_text)
                        _appliquer_style_run(run, None, base_style)
            continue

        if type_bloc.startswith("heading"):
            try:
                level = int(type_bloc.split("_")[-1])
            except Exception:
                level = 1
            p = document.add_heading(level=level)
        else:
            p = document.add_paragraph()

        for run_data in bloc.get("runs", []):
            run = p.add_run(run_data.get("text", ""))
            _appliquer_style_run(run, run_data.get("style"), base_style)

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output

